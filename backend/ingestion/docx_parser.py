from dataclasses import dataclass, field
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


@dataclass
class ParsedBlock:
    text: str
    page_num: int = 0
    font_size: float = 12.0
    is_bold: bool = False
    is_table: bool = False
    style_name: str = "Normal"
    bbox: tuple = field(default_factory=lambda: (0, 0, 0, 0))


def parse_docx(path: Path) -> tuple[list[ParsedBlock], int]:
    doc = Document(str(path))
    blocks: list[ParsedBlock] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else "Normal"
        is_bold = any(run.bold for run in para.runs if run.bold is not None)
        font_size = 12.0
        for run in para.runs:
            if run.font.size:
                font_size = run.font.size.pt
                break

        blocks.append(ParsedBlock(
            text=text,
            font_size=font_size,
            is_bold=is_bold,
            style_name=style,
        ))

    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        table_text = "\n".join(rows)
        if table_text.strip():
            blocks.append(ParsedBlock(
                text=table_text,
                is_table=True,
                style_name="Table",
            ))

    # DOCX has no reliable page count without rendering; estimate from content length
    total_chars = sum(len(b.text) for b in blocks)
    estimated_pages = max(1, total_chars // 3000)
    return blocks, estimated_pages
